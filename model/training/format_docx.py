"""
Форматирование результатов распознавания маркеров в красивый docx.
Читает JSON-файлы из каталога и text.csv, строит отчёт.

Использование:
  python format_docx.py <имя_каталога>
  python format_docx.py dkatja
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
import re
import json
import csv
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# --- Описание маркеров ---
MARKER_DESC = {
    '1': 'Парная замена ж/х',
    '2': 'Парная замена н/п, и/н',
    '3': 'Парная замена в/з',
    '4': 'Парная замена т/г',
    '5': 'Парная замена р/ь/б',
    '6': 'Парная замена о/ю',
    '7': 'Замена по звонкости/глухости',
    '8': 'Парная замена х/к',
    '9': 'Пропуск согласной на стечении',
    '0': 'Нет ошибки',
    'A': 'Удлинение гласной с пропуском согласной',
    'B': 'Пропуск гласной между согласными',
    'C': 'Произвольная замена',
    'D': 'Перестановка гласных',
    'E': 'Перестановка согласных',
    'F': 'Лишний звук в конце слова',
    'G': 'Смягчение окончания',
    'H': 'Пропуск с удвоением',
    'J': 'Повтор через паузу / лишний префикс',
    'K': 'Заикание',
    'L': 'Удлинение согласной с пропуском гласной',
}

# --- Цвета для маркеров ---
MARKER_COLOR = {
    '1': RGBColor(0x8B, 0x00, 0x00),
    '2': RGBColor(0x8B, 0x00, 0x00),
    '3': RGBColor(0x8B, 0x00, 0x00),
    '4': RGBColor(0x8B, 0x00, 0x00),
    '5': RGBColor(0x8B, 0x00, 0x00),
    '6': RGBColor(0x8B, 0x00, 0x00),
    '7': RGBColor(0xCC, 0x66, 0x00),
    '8': RGBColor(0x8B, 0x00, 0x00),
    '9': RGBColor(0x00, 0x66, 0xCC),
    'A': RGBColor(0x99, 0x33, 0x99),
    'B': RGBColor(0x00, 0x66, 0xCC),
    'C': RGBColor(0x66, 0x66, 0x66),
    'D': RGBColor(0x00, 0x88, 0x00),
    'E': RGBColor(0x00, 0x88, 0x00),
    'F': RGBColor(0xCC, 0x00, 0x66),
    'G': RGBColor(0xCC, 0x66, 0x00),
    'H': RGBColor(0x00, 0x66, 0xCC),
    'J': RGBColor(0x99, 0x33, 0x99),
    'K': RGBColor(0x99, 0x33, 0x99),
    'L': RGBColor(0x99, 0x33, 0x99),
}


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_text(cell, text, font_name='Times New Roman', font_size=11,
                       bold=False, italic=False, color=None, alignment=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    lines = text.split('\n')
    for li, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
        if li < len(lines) - 1:
            run.add_break()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def parse_json_to_words(data):
    """
    Из JSON [[t1, f2, mark], ...] собрать слова.
    Возвращает список словарей:
      {'ref': 'деревне', 'hyp': 'деревни', 'marks': '6'}
    """
    words = []
    cur_ref = ''
    cur_hyp = ''
    cur_marks = ''

    for row in data:
        t1, f2, mark = row[0], row[1], row[2]
        if t1 == '|' and f2 == '|':
            # Граница слова — сохраняем текущее слово
            if cur_ref or cur_hyp:
                words.append({
                    'ref': cur_ref,
                    'hyp': cur_hyp,
                    'marks': cur_marks
                })
            cur_ref = ''
            cur_hyp = ''
            cur_marks = ''
        elif t1 == '|' and f2 != '|':
            # f2 содержит лишний звук на границе слова
            # Добавляем к гипотезе предыдущего слова
            cur_hyp += f2
            if mark != '0':
                cur_marks += mark
        elif t1 != '|' and f2 == '|':
            # Пропуск звука
            cur_ref += t1
            if mark != '0':
                cur_marks += mark
        else:
            cur_ref += t1
            cur_hyp += f2
            if mark != '0':
                cur_marks += mark

    # Последнее слово
    if cur_ref or cur_hyp:
        words.append({
            'ref': cur_ref,
            'hyp': cur_hyp,
            'marks': cur_marks
        })

    return words


def build_sentence_row(words):
    """
    Из списка слов собрать строку таблицы:
    - etalon: все слова эталона через пробел
    - produced: все слова гипотезы через пробел
    - markers_text: построчно "Слово коды" только для слов с ошибками
    """
    etalon = ' '.join(w['ref'] for w in words if w['ref'])
    produced = ' '.join(w['hyp'] for w in words if w['hyp'])

    marker_lines = []
    for w in words:
        if w['marks']:
            # Показываем слово гипотезы + коды маркеров
            word_label = w['hyp'] if w['hyp'] else w['ref']
            marker_lines.append(f"{word_label} {w['marks']}")

    markers_text = '\n'.join(marker_lines)
    return etalon, produced, markers_text


def main():
    # --- Аргументы ---
    if len(sys.argv) < 2:
        print('Использование: python format_docx.py <имя_каталога>')
        print('Пример: python format_docx.py dkatja')
        sys.exit(1)

    src_dir = sys.argv[1]
    dir_name = os.path.basename(src_dir)

    # --- Читаем text.csv ---
    reader = csv.reader(open('text.csv', encoding='utf-8'), delimiter='\t', quotechar='|')
    texts = [row[0].strip() for row in reader]

    # --- Находим JSON-файлы ---
    json_files = sorted([
        f for f in os.listdir(src_dir)
        if f.endswith('.json')
    ])

    # --- Собираем данные по предложениям ---
    sentences = []
    for jf in json_files:
        match = re.search(r'(\d+)\.json$', jf)
        if not match:
            continue
        idx = int(match.group(1))
        ref_text = texts[idx - 1] if idx - 1 < len(texts) else ''

        with open(os.path.join(src_dir, jf), encoding='utf-8') as f:
            data = json.load(f)

        words = parse_json_to_words(data)
        _, produced, markers_text = build_sentence_row(words)
        sentences.append({
            'idx': idx,
            'ref_text': ref_text,
            'etalon': ref_text,  # берём из text.csv, не из фонем
            'produced': produced,
            'markers': markers_text,
            'words': words,  # для пословной раскраски
        })

    print(f'Каталог: {dir_name}, предложений: {len(sentences)}')

    # --- Создаём документ ---
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

    # --- Заголовок ---
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('Результаты анализа речи')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f'Каталог: {dir_name}')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # --- Основная таблица ---
    num_rows = len(sentences) + 1  # +1 для заголовка
    table = doc.add_table(rows=num_rows, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    widths = [Cm(1), Cm(4.5), Cm(5.5), Cm(3), Cm(4)]
    header_color = '1A1A5E'

    headers = ['№', 'Эталон', 'Произнесено', 'Маркеры', 'Расшифровка']
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        set_cell_shading(cell, header_color)
        add_formatted_text(cell, h, font_size=11, bold=True,
                          color=RGBColor(0xFF, 0xFF, 0xFF),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        cell.width = widths[ci]

    # --- Заполняем данные ---
    for ri, sent in enumerate(sentences):
        row_idx = ri + 1
        bg = 'F0F4FA' if row_idx % 2 == 0 else 'FFFFFF'

        # №
        cell = table.rows[row_idx].cells[0]
        set_cell_shading(cell, bg)
        add_formatted_text(cell, str(sent['idx']), font_size=10,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        cell.width = widths[0]

        # Эталон
        cell = table.rows[row_idx].cells[1]
        set_cell_shading(cell, bg)
        add_formatted_text(cell, sent['etalon'], font_size=11)
        cell.width = widths[1]

        # Произнесено — пословная раскраска
        cell = table.rows[row_idx].cells[2]
        set_cell_shading(cell, bg)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for wi, w in enumerate(sent['words']):
            word_text = w['hyp'] if w['hyp'] else w['ref']
            if not word_text:
                continue
            if wi > 0:
                sp = p.add_run(' ')
                sp.font.name = 'Times New Roman'
                sp.font.size = Pt(11)
            run = p.add_run(word_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            if w['marks']:
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            else:
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        if not sent['words']:
            run = p.add_run('-')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        cell.width = widths[2]

        # Маркеры
        cell = table.rows[row_idx].cells[3]
        set_cell_shading(cell, bg)
        if sent['markers']:
            add_formatted_text(cell, sent['markers'], font_size=10, bold=True,
                              color=RGBColor(0x8B, 0x00, 0x00),
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            add_formatted_text(cell, '\u2713', font_size=12, bold=True,
                              color=RGBColor(0x00, 0x80, 0x00),
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)
        cell.width = widths[3]

        # Расшифровка
        cell = table.rows[row_idx].cells[4]
        set_cell_shading(cell, bg)
        if sent['markers']:
            codes = set()
            for ch in sent['markers']:
                if ch in MARKER_DESC:
                    codes.add(ch)
            desc_lines = [f'{c} \u2014 {MARKER_DESC[c]}' for c in sorted(codes)]
            if desc_lines:
                cell.text = ''
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for di, line in enumerate(desc_lines):
                    code = line[0]
                    clr = MARKER_COLOR.get(code, RGBColor(0x66, 0x66, 0x66))
                    run = p.add_run(line)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = clr
                    if di < len(desc_lines) - 1:
                        run.add_break()
        else:
            add_formatted_text(cell, 'Без ошибок', font_size=9, italic=True,
                              color=RGBColor(0x00, 0x80, 0x00))
        cell.width = widths[4]

    doc.add_paragraph()

    # --- Статистика ---
    total = len(sentences)
    no_errors = sum(1 for s in sentences if not s['markers'])
    with_errors = total - no_errors

    # Подсчёт маркеров
    all_codes = {}
    for s in sentences:
        for ch in s['markers']:
            if ch in MARKER_DESC:
                all_codes[ch] = all_codes.get(ch, 0) + 1

    stat_heading = doc.add_paragraph()
    run = stat_heading.add_run('Статистика')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)

    stat_p = doc.add_paragraph()
    stat_p.paragraph_format.space_after = Pt(4)
    for text, val in [
        ('Всего предложений: ', str(total)),
        ('Без ошибок: ', str(no_errors)),
        ('С ошибками: ', str(with_errors)),
    ]:
        run = stat_p.add_run(text)
        run.font.size = Pt(11)
        run = stat_p.add_run(val)
        run.font.size = Pt(11)
        run.font.bold = True
        run.add_break()

    if all_codes:
        doc.add_paragraph()
        freq_heading = doc.add_paragraph()
        run = freq_heading.add_run('Частота маркеров')
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)

        ft = doc.add_table(rows=len(all_codes) + 1, cols=3)
        ft.style = 'Table Grid'
        ft.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ci, h in enumerate(['Код', 'Описание', 'Кол-во']):
            cell = ft.rows[0].cells[ci]
            set_cell_shading(cell, header_color)
            add_formatted_text(cell, h, font_size=10, bold=True,
                              color=RGBColor(0xFF, 0xFF, 0xFF),
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)

        for li, (code, count) in enumerate(sorted(all_codes.items(),
                                                   key=lambda x: -x[1])):
            ri = li + 1
            bg = 'F0F4FA' if ri % 2 == 0 else 'FFFFFF'
            clr = MARKER_COLOR.get(code, RGBColor(0x66, 0x66, 0x66))

            cell = ft.rows[ri].cells[0]
            set_cell_shading(cell, bg)
            add_formatted_text(cell, code, font_size=11, bold=True,
                              color=clr, alignment=WD_ALIGN_PARAGRAPH.CENTER)

            cell = ft.rows[ri].cells[1]
            set_cell_shading(cell, bg)
            add_formatted_text(cell, MARKER_DESC.get(code, '?'), font_size=10,
                              color=clr)

            cell = ft.rows[ri].cells[2]
            set_cell_shading(cell, bg)
            add_formatted_text(cell, str(count), font_size=11, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # --- Легенда ---
    legend_heading = doc.add_paragraph()
    run = legend_heading.add_run('Легенда маркеров')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)

    legend_items = sorted(MARKER_DESC.items(), key=lambda x: x[0])
    lt = doc.add_table(rows=len(legend_items) + 1, cols=2)
    lt.style = 'Table Grid'
    lt.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, h in enumerate(['Код', 'Описание']):
        cell = lt.rows[0].cells[ci]
        set_cell_shading(cell, header_color)
        add_formatted_text(cell, h, font_size=10, bold=True,
                          color=RGBColor(0xFF, 0xFF, 0xFF),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for li, (code, desc) in enumerate(legend_items):
        ri = li + 1
        bg = 'F0F4FA' if ri % 2 == 0 else 'FFFFFF'
        clr = MARKER_COLOR.get(code, RGBColor(0x66, 0x66, 0x66))

        cell = lt.rows[ri].cells[0]
        set_cell_shading(cell, bg)
        add_formatted_text(cell, code, font_size=11, bold=True,
                          color=clr, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        cell = lt.rows[ri].cells[1]
        set_cell_shading(cell, bg)
        add_formatted_text(cell, desc, font_size=10, color=clr)

    # --- Сохраняем ---
    if len(sys.argv) >= 3:
        out_dir = sys.argv[2]
        os.makedirs(out_dir, exist_ok=True)
        output_path = os.path.join(out_dir, f'{dir_name}_report.docx')
    else:
        output_path = os.path.join(src_dir, f'{dir_name}_report.docx')
    doc.save(output_path)
    print(f'Сохранено: {output_path}')


if __name__ == '__main__':
    main()
